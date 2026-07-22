# -*- coding: utf-8 -*-
"""roteiro_doc.py — exporta o roteiro.txt para .docx (e .pdf, se possível).

Regra do usuário: o roteiro SEMPRE deve ficar disponível em DOCS (.docx) ou PDF para
entregar à equipe. O pipeline chama exportar() automaticamente depois que o roteiro fica
pronto/validado; também roda standalone:

    py -3 roteiro_doc.py "<pasta_do_projeto>"

Sempre gera o .docx (python-docx). O .pdf é "bônus" via Word COM (pywin32) — se o Word
não estiver instalado ou a conversão falhar, segue só com o .docx (que já satisfaz a regra).
Regenera só quando o roteiro.txt é mais novo que o .docx (idempotente, mas sempre atual).
"""

from __future__ import annotations  # tipos 'X | None' em Python 3.9 (macOS)

import sys
from pathlib import Path

# Fontes Unicode (Windows) p/ o PDF sair com travessões/aspas/acentos corretos.
# Ordem de preferência: (regular, bold). A 1ª que existir é usada.
_FONTES = (
    (r"C:\Windows\Fonts\georgia.ttf", r"C:\Windows\Fonts\georgiab.ttf"),
    (r"C:\Windows\Fonts\times.ttf",   r"C:\Windows\Fonts\timesbd.ttf"),
    (r"C:\Windows\Fonts\arial.ttf",   r"C:\Windows\Fonts\arialbd.ttf"),
)


def _partes(texto):
    """1ª linha não-vazia = título; demais blocos separados por linha em branco = parágrafos."""
    linhas = texto.replace("\r\n", "\n").split("\n")
    titulo = ""
    i = 0
    for i, l in enumerate(linhas):
        if l.strip():
            titulo = l.strip()
            break
    resto = "\n".join(linhas[i + 1:])
    paras = [p.strip() for p in resto.split("\n\n")]
    paras = [p for p in paras if p]
    return titulo, paras


def _gerar_docx(titulo, paras, saida):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # margens confortáveis para leitura/impressão
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1.0)
        sec.top_margin = sec.bottom_margin = Inches(1.0)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Georgia"
    estilo.font.size = Pt(12)

    t = doc.add_heading(titulo, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in paras:
        par = doc.add_paragraph(p)
        par.paragraph_format.space_after = Pt(10)
        par.paragraph_format.line_spacing = 1.3
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(str(saida))


def _achar_fonte():
    """Devolve (regular, bold) — caminhos de TTF existentes; bold cai p/ regular se faltar."""
    import os
    for reg, bold in _FONTES:
        if os.path.isfile(reg):
            return reg, (bold if os.path.isfile(bold) else reg)
    return None, None


def _gerar_pdf_fpdf(titulo, paras, pdf_path, log):
    """PDF nativo via fpdf2 (sem depender de Word/LibreOffice). Fonte Unicode do Windows."""
    try:
        from fpdf import FPDF
    except Exception:
        return False
    reg, bold = _achar_fonte()
    if not reg:
        log("    (PDF fpdf2 pulado: nenhuma fonte TTF encontrada em C:\\Windows\\Fonts.)")
        return False
    try:
        pdf = FPDF(format="A4", unit="mm")
        pdf.set_margins(20, 20, 20)
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_font("roteiro", "", reg)
        pdf.add_font("roteiro", "B", bold)
        pdf.add_page()
        # título
        pdf.set_font("roteiro", "B", 17)
        pdf.multi_cell(0, 9, titulo, align="C")
        pdf.ln(5)
        # corpo justificado
        pdf.set_font("roteiro", "", 12)
        for p in paras:
            pdf.multi_cell(0, 7, p, align="J")
            pdf.ln(3)
        pdf.output(str(pdf_path))
        return True
    except Exception as e:  # noqa: BLE001
        log("    (PDF fpdf2 falhou: %s)" % e)
        return False


def _gerar_pdf_word(docx_path, pdf_path, log):
    """Reserva: converte o .docx em .pdf usando o Word (COM). Silencioso se indisponível."""
    try:
        import pythoncom
        import win32com.client as win32
    except Exception:
        log("    (PDF pulado: pywin32 indisponível — .docx já entregue.)")
        return False
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(docx_path))
        d.SaveAs(str(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        d.Close(False)
        return True
    except Exception as e:  # noqa: BLE001
        log("    (PDF pulado: Word indisponível/erro: %s — .docx já entregue.)" % e)
        return False
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def exportar(roteiro_txt, log=print, fazer_pdf=True):
    """Gera <pasta>/roteiro.docx (+ .pdf) a partir de roteiro.txt. Retorna (docx, pdf|None)."""
    roteiro_txt = Path(roteiro_txt)
    if not (roteiro_txt.exists() and roteiro_txt.stat().st_size > 0):
        return None, None
    docx_path = roteiro_txt.with_suffix(".docx")
    pdf_path = roteiro_txt.with_suffix(".pdf")

    titulo, paras = _partes(roteiro_txt.read_text(encoding="utf-8", errors="replace"))

    # idempotência: se o .docx está atualizado, só (re)gera o PDF caso esteja faltando.
    if docx_path.exists() and docx_path.stat().st_mtime >= roteiro_txt.stat().st_mtime:
        if fazer_pdf and not pdf_path.exists():
            ok = _gerar_pdf_fpdf(titulo, paras, pdf_path, log) or _gerar_pdf_word(docx_path, pdf_path, log)
            if ok:
                log("    ✓ roteiro.pdf gerado.")
            return docx_path, (pdf_path if ok else None)
        log("    roteiro.docx/pdf já estão atualizados — exportação pulada.")
        return docx_path, (pdf_path if pdf_path.exists() else None)

    _gerar_docx(titulo, paras, docx_path)
    log("    ✓ roteiro.docx gerado (%d parágrafos)." % len(paras))

    pdf_ok = False
    if fazer_pdf:
        # fpdf2 primeiro (não depende de Word/LibreOffice); Word COM como reserva.
        pdf_ok = _gerar_pdf_fpdf(titulo, paras, pdf_path, log)
        if not pdf_ok:
            pdf_ok = _gerar_pdf_word(docx_path, pdf_path, log)
        if pdf_ok:
            log("    ✓ roteiro.pdf gerado.")
    return docx_path, (pdf_path if pdf_ok else None)


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    if len(sys.argv) < 2:
        print('uso: py -3 roteiro_doc.py "<pasta_do_projeto>"')
        return
    pasta = Path(sys.argv[1])
    txt = pasta / "roteiro.txt" if pasta.is_dir() else pasta
    docx_path, pdf_path = exportar(txt)
    if docx_path:
        print("DOCX:", docx_path)
        print("PDF :", pdf_path or "(não gerado)")
    else:
        print("roteiro.txt não encontrado em", pasta)


if __name__ == "__main__":
    main()
